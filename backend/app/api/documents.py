"""
Documents API - LangChain RAG Pipeline
Upload → chunk → embed → ChromaDB → retrieve → LLM answer
"""
import os
import uuid
from typing import List
from fastapi import APIRouter, Depends, UploadFile, File, HTTPException
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from pydantic import BaseModel, ConfigDict

from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import FakeEmbeddings
from langchain_core.documents import Document as LangchainDocument
from app.agents.claude_agent import run_chat

from app.db.session import get_db
from app.db.models import Document
from app.utils.config import settings
from app.api.deps import get_current_user

router = APIRouter()


# ── Embeddings & Vector Store ─────────────────────────────────────
# Using FakeEmbeddings for Railway deployment - replace with
# a real embeddings API (OpenAI/Cohere) for production use
def get_embeddings():
    return FakeEmbeddings(size=384)


def get_chroma_client():
    import chromadb
    return chromadb.HttpClient(host=settings.CHROMA_HOST, port=settings.CHROMA_PORT)


# ── Upload Document ───────────────────────────────────────────────
@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    current_user=Depends(get_current_user),
):
    """Upload and index a document into ChromaDB for RAG queries."""
    allowed = {".pdf", ".txt", ".md", ".docx", ".csv"}
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in allowed:
        raise HTTPException(400, f"File type {ext} not supported. Allowed: {allowed}")

    # Save file
    file_id = str(uuid.uuid4())
    os.makedirs(settings.UPLOAD_DIR, exist_ok=True)
    file_path = os.path.join(settings.UPLOAD_DIR, f"{file_id}{ext}")
    content = await file.read()
    with open(file_path, "wb") as f:
        f.write(content)

    # Load document with LangChain
    if ext == ".pdf":
        loader = PyPDFLoader(file_path)
    elif ext in (".txt", ".md", ".csv"):
        loader = TextLoader(file_path, encoding="utf-8")
    elif ext == ".docx":
        from docx import Document as DocxDocument
        docx_file = DocxDocument(file_path)
        text = "\n".join([p.text for p in docx_file.paragraphs if p.text.strip()])
        docs = [LangchainDocument(page_content=text, metadata={"source": file_path})]
    else:
        raise HTTPException(400, "Unsupported file type")

    if ext != ".docx":
        docs = loader.load()

    # Split into chunks
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.CHUNK_SIZE,
        chunk_overlap=settings.CHUNK_OVERLAP,
    )
    chunks = splitter.split_documents(docs)

    # Generate collection name per user-document
    collection_name = f"user_{current_user.id[:8]}_{file_id[:8]}"

    # Embed and store in ChromaDB
    embeddings = get_embeddings()
    Chroma.from_documents(
        documents=chunks,
        embedding=embeddings,
        client=get_chroma_client(),
        collection_name=collection_name,
    )

    # Save metadata
    doc = Document(
        user_id=current_user.id,
        file_name=file.filename,
        file_path=file_path,
        file_size=len(content),
        mime_type=file.content_type,
        chunk_count=len(chunks),
        is_indexed=True,
        collection_name=collection_name,
    )
    db.add(doc)
    await db.commit()
    await db.refresh(doc)

    return {
        "id": doc.id,
        "file_name": doc.file_name,
        "chunk_count": doc.chunk_count,
        "status": "indexed",
    }


# ── Query Documents (RAG) ─────────────────────────────────────────
class QueryRequest(BaseModel):
    model_config = ConfigDict(protected_namespaces=())

    question: str
    document_ids: List[str] = []
    model_name: str = "llama3"


@router.post("/query")
async def query_documents(
    payload: QueryRequest,
    db: AsyncSession = Depends(get_db),
    current_user=Depends(get_current_user),
):
    """RAG query: retrieve relevant chunks + generate answer with LLM."""
    # Get user's documents
    if payload.document_ids:
        result = await db.execute(
            select(Document).where(
                Document.id.in_(payload.document_ids),
                Document.user_id == current_user.id,
                Document.is_indexed == True,
            )
        )
    else:
        result = await db.execute(
            select(Document).where(
                Document.user_id == current_user.id,
                Document.is_indexed == True,
            )
        )
    docs = result.scalars().all()
    if not docs:
        raise HTTPException(404, "No indexed documents found")

    embeddings = get_embeddings()
    chroma_client = get_chroma_client()

    # Query each document's collection and aggregate results
    combined_retriever_docs = []
    for doc in docs:
        try:
            vectorstore = Chroma(
                client=chroma_client,
                collection_name=doc.collection_name,
                embedding_function=embeddings,
            )
            retrieved = vectorstore.similarity_search_with_score(payload.question, k=2)
            for chunk, score in retrieved:
                chunk.metadata["source_file"] = doc.file_name
                chunk.metadata["similarity_score"] = float(score)
                combined_retriever_docs.append((chunk, score))
        except Exception:
            continue

    if not combined_retriever_docs:
        raise HTTPException(404, "No relevant chunks found")

    # Sort by similarity, take top-k
    combined_retriever_docs.sort(key=lambda x: x[1])
    top_docs = [d for d, _ in combined_retriever_docs[:settings.RAG_TOP_K]]

    # Build context
    context = "\n\n---\n\n".join([d.page_content for d in top_docs])

    # Generate answer with Claude
    prompt = f"""Answer the following question based ONLY on the provided context.
If the answer is not in the context, say "I cannot find this information in the uploaded documents."

Context:
{context}

Question: {payload.question}

Answer:"""

    model_name = payload.model_name if hasattr(payload, 'model_name') and payload.model_name else "claude-sonnet"
    answer = run_chat(
        agent_name="RAG Assistant",
        system_prompt="You are a document Q&A assistant. Answer questions based only on the provided context.",
        model_name=model_name,
        temperature=0.3,
        max_tokens=2048,
        input_text=prompt,
    )

    sources = [
        {
            "file_name": d.metadata.get("source_file", "unknown"),
            "page": d.metadata.get("page", 0),
            "similarity": d.metadata.get("similarity_score", 0),
            "excerpt": d.page_content[:200],
        }
        for d in top_docs
    ]

    return {"answer": answer, "sources": sources, "chunks_retrieved": len(top_docs)}


# ── List Documents ────────────────────────────────────────────────
@router.get("/")
async def list_documents(
    db: AsyncSession = Depends(get_db),
    current_user=Depends(get_current_user),
):
    result = await db.execute(
        select(Document)
        .where(Document.user_id == current_user.id)
        .order_by(Document.uploaded_at.desc())
    )
    docs = result.scalars().all()
    return [
        {
            "id": d.id,
            "file_name": d.file_name,
            "file_size": d.file_size,
            "chunk_count": d.chunk_count,
            "is_indexed": d.is_indexed,
            "uploaded_at": d.uploaded_at,
        }
        for d in docs
    ]


@router.delete("/{document_id}")
async def delete_document(
    document_id: str,
    db: AsyncSession = Depends(get_db),
    current_user=Depends(get_current_user),
):
    result = await db.execute(
        select(Document).where(Document.id == document_id, Document.user_id == current_user.id)
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(404, "Document not found")

    # Remove from ChromaDB
    try:
        chroma_client = get_chroma_client()
        chroma_client.delete_collection(doc.collection_name)
    except Exception:
        pass

    # Remove file
    try:
        os.remove(doc.file_path)
    except Exception:
        pass

    await db.delete(doc)
    await db.commit()
    return {"status": "deleted"}
